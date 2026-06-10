#!/usr/bin/env python3
"""Parse the "Open Research Across Disciplines" Google Doc into the disciplines spreadsheet.

This script extracts the contents of the canonical hackathon document
(a .docx in Google Drive) and rewrites the three tabs of the disciplines
Google Sheet (Fields / Disciplines / Resources) consumed by
`parse_disciplines_to_json.py`.

Use it when a new hackathon edition of the G-Doc replaces the previous
content wholesale. It is intentionally destructive: it clears each tab
before writing. Save a backup first if you want one.

Expected document structure (Word/Google Docs styles):

    Heading 1   = Field name (e.g. "Natural Sciences")
                  Headings "Introduction" and "Contents" are ignored.
    Heading 2   = Discipline name within a field (e.g. "Chemistry").
                  If a field has resources but no Heading 2, a synthetic
                  discipline matching the field name is created.
    Heading 3   = Section label ("Examples of open research practices",
                  "Table of resources"). Only the heading text is used
                  to decide what to do with subsequent paragraphs.
    Paragraphs under an "Examples..." Heading 3 become the discipline's
    Examples text.
    Tables under a "Table of resources" Heading 3 become resource rows.
    Resource tables are 2-column (category | resource) or 1-column
    (resource only). Cells with multiple URLs are split into one
    resource per URL, sharing the same title.

Categories are taken from the first colon-delimited token of the
category cell ("Open Data: COS" → "Open Data").

Usage:
    pip install python-docx
    # Either point at a local .docx ...
    python scripts/parse_disciplines_gdoc.py path/to/source.docx
    # ... or pull from Drive via gws (requires the gws CLI configured)
    python scripts/parse_disciplines_gdoc.py --drive-id <FILE_ID>

    # Add --push to overwrite the disciplines sheet after parsing.
    python scripts/parse_disciplines_gdoc.py --drive-id <FILE_ID> --push

    # Inspect what *would* be written without touching the sheet:
    python scripts/parse_disciplines_gdoc.py path/to/source.docx --json /tmp/parsed.json

Sheet IDs are constants below; update them if the canonical sheet moves.
"""

import argparse
import json
import os
import re
import subprocess
import sys
import tempfile

from docx import Document

# Google Sheet that drives the FORRT disciplines Hugo page.
DISCIPLINES_SHEET_ID = "1mSlduu86_nE1sY1gXobw3Pp1vI73B_0iHBsJqjtsJU4"

# Fields whose Heading 1 we never want to import — they are document-level
# front matter, not topical fields.
SKIP_FIELDS = {"Introduction", "Contents"}

# Fields that should appear on the public page (Show = TRUE). Others are
# kept in the sheet but hidden by the Hugo build. Adjust if scope changes.
SHOW_FIELDS = {
    "Natural Sciences",
    "Life, Medical and Health Sciences",
    "Social Sciences",
    "Humanities",
    "Engineering and Technology",
    "Meta-research",
    "Methodologies",
}

# A URL token: lazy match that stops at whitespace, closing brackets, or
# the start of the next URL — so two URLs concatenated without a separator
# (a common artifact when Google Docs hyperlinks are stripped) split cleanly.
URL_RE = re.compile(r"https?://[^\s)<>]+?(?=https?://|\s|$|[)<>])")


# ---------- docx parsing ---------------------------------------------------


def _clean_title(s: str) -> str:
    """Collapse whitespace and trim trailing punctuation, preserving 'e.g.'-style endings."""
    t = re.sub(r"\s+", " ", s).strip()
    if t.endswith(".") and not re.search(r"\b[a-z]\.[a-z]\.$", t):
        t = t[:-1]
    return t.strip(" \t").rstrip(":,;").strip()


def _resource_cell_texts(cells):
    """Return (category_text, resource_text) from a 1- or 2-column resource row."""
    if len(cells) == 1:
        return "", cells[0].text
    return cells[0].text, cells[1].text


def parse_cell_resources(cat_txt: str, res_txt: str):
    """Convert one resource cell into a list of {title, link, category} dicts.

    Cells frequently contain a short descriptor followed by one *or more* URLs;
    we emit one resource per URL, all sharing the descriptor as the title. A
    cell with no URLs becomes a single resource with an empty link.
    """
    category = (cat_txt.strip().split(":", 1)[0].strip()) if cat_txt.strip() else ""
    text = res_txt.strip()

    spans, urls = [], []
    for m in URL_RE.finditer(text):
        spans.append((m.start(), m.end()))
        urls.append(m.group(0).rstrip(".,;"))

    if not urls:
        return [{"title": _clean_title(text), "link": "", "category": category}]

    # Title prefix = the non-URL text before the first URL. If empty
    # (cell starts with a URL), concatenate any inter-URL text as a fallback.
    prefix = _clean_title(text[: spans[0][0]])
    if not prefix:
        parts, last = [], 0
        for s, e in spans:
            parts.append(text[last:s])
            last = e
        parts.append(text[last:])
        prefix = _clean_title(" ".join(parts))

    return [{"title": prefix, "link": url, "category": category} for url in urls]


def parse_docx(path: str):
    """Walk a .docx and return {"fields": [...]} in document order."""
    doc = Document(path)
    fields = []
    current_field = None
    current_disc = None
    last_h3 = None
    skip_field = False
    title_table_consumed = False

    body = doc.element.body
    p_idx = 0
    t_idx = 0

    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = doc.paragraphs[p_idx]
            p_idx += 1
            style = p.style.name if p.style else ""
            txt = p.text.strip()
            if not txt:
                continue
            if style == "Heading 1":
                if txt in SKIP_FIELDS:
                    skip_field = True
                    continue
                skip_field = False
                current_field = {
                    "name": txt, "summary": "",
                    "disciplines": [], "field_examples": "",
                }
                fields.append(current_field)
                current_disc = None
                last_h3 = None
            elif skip_field:
                continue
            elif style == "Heading 2":
                current_disc = {"name": txt, "examples": "", "resources": []}
                current_field["disciplines"].append(current_disc)
                last_h3 = None
            elif style == "Heading 3":
                last_h3 = txt
            else:
                if last_h3 and "Examples" in last_h3:
                    target = current_disc if current_disc is not None else current_field
                    key = "examples" if current_disc is not None else "field_examples"
                    existing = target.get(key, "")
                    target[key] = (existing + ("\n" if existing else "") + txt)
        elif tag == "tbl":
            tbl = doc.tables[t_idx]
            t_idx += 1
            # First table in the body is the cover page (authors, title, etc.)
            if not title_table_consumed:
                title_table_consumed = True
                continue
            if skip_field or current_field is None or len(tbl.rows) < 2:
                continue
            if current_disc is None:
                # Field-level table — create a synthetic discipline named after the field
                current_disc = {
                    "name": current_field["name"],
                    "examples": current_field.get("field_examples", ""),
                    "resources": [],
                }
                current_field["disciplines"].append(current_disc)
            for row in tbl.rows[1:]:
                cells = list(row.cells)
                cat_txt, res_txt = _resource_cell_texts(cells)
                if not res_txt.strip():
                    continue
                current_disc["resources"].extend(parse_cell_resources(cat_txt, res_txt))
            last_h3 = None
        # other body children (sdt, sectPr) are ignored

    return {"fields": fields}


# ---------- sheet I/O ------------------------------------------------------


def _gws(*args, body=None):
    """Invoke the gws CLI, strip its keyring banner, and return parsed JSON output."""
    cmd = ["gws", *args, "--format", "json"]
    if body is not None:
        cmd += ["--json", json.dumps(body)]
    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        sys.exit(f"gws failed: {' '.join(cmd)}\n{res.stderr}")
    out = res.stdout.strip().split("\n")
    if out and out[0].startswith("Using"):
        out = out[1:]
    return json.loads("\n".join(out)) if out else {}


def download_docx_from_drive(file_id: str, dest_path: str) -> str:
    """Download a Google Drive .docx by file ID via gws and return its path."""
    print(f"Downloading docx (id={file_id}) → {dest_path}")
    _gws(
        "drive", "files", "get",
        "--params", json.dumps({"fileId": file_id, "alt": "media"}),
        "--output", dest_path,
    )
    return dest_path


def build_sheet_payloads(data: dict):
    """Convert parsed JSON into the three header+rows blocks expected by the sheet."""
    fields_rows = [["Name", "Summary", "Show"]]
    for f in data["fields"]:
        show = "TRUE" if f["name"] in SHOW_FIELDS else "FALSE"
        fields_rows.append([f["name"], (f.get("summary") or "").strip(), show])

    disc_rows = [["Field", "Discipline", "Examples"]]
    for f in data["fields"]:
        for d in f["disciplines"]:
            disc_rows.append([f["name"], d["name"], (d.get("examples") or "").strip()])

    res_rows = [["Discipline", "Title", "Link", "Category"]]
    for f in data["fields"]:
        for d in f["disciplines"]:
            for r in d["resources"]:
                res_rows.append([d["name"], r["title"], r.get("link", ""), r.get("category", "")])

    return fields_rows, disc_rows, res_rows


def push_to_sheet(spreadsheet_id: str, payloads):
    """Clear and rewrite the Fields / Disciplines / Resources tabs."""
    tabs = [("Fields", payloads[0]), ("Disciplines", payloads[1]), ("Resources", payloads[2])]
    for tab, rows in tabs:
        print(f"Clearing {tab}…")
        _gws(
            "sheets", "spreadsheets", "values", "clear",
            "--params", json.dumps({"spreadsheetId": spreadsheet_id, "range": tab}),
        )
        print(f"Writing {len(rows)} rows to {tab}…")
        result = _gws(
            "sheets", "spreadsheets", "values", "update",
            "--params", json.dumps({
                "spreadsheetId": spreadsheet_id, "range": tab, "valueInputOption": "RAW",
            }),
            body={"values": rows},
        )
        print(f"  updatedRange={result.get('updatedRange')}  rows={result.get('updatedRows')}")


# ---------- CLI ------------------------------------------------------------


def print_stats(data: dict) -> None:
    total_disc = total_res = 0
    for f in data["fields"]:
        n_disc = len(f["disciplines"])
        n_res = sum(len(d["resources"]) for d in f["disciplines"])
        total_disc += n_disc
        total_res += n_res
        print(f"  {f['name']:<55} disciplines={n_disc:2d}  resources={n_res:4d}")
    print(f"Total: {len(data['fields'])} fields, {total_disc} disciplines, {total_res} resources")


def main():
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    src = parser.add_mutually_exclusive_group(required=True)
    src.add_argument("docx_path", nargs="?", help="Local .docx file to parse")
    src.add_argument("--drive-id", help="Google Drive file ID of the .docx")
    parser.add_argument(
        "--json", default=None,
        help="Optional path to write parsed JSON (for review). Default: skip.",
    )
    parser.add_argument(
        "--push", action="store_true",
        help="Clear and rewrite the disciplines spreadsheet from the parsed content.",
    )
    parser.add_argument(
        "--spreadsheet-id", default=DISCIPLINES_SHEET_ID,
        help="Override the target spreadsheet ID.",
    )
    args = parser.parse_args()

    if args.drive_id:
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.close()
        docx_path = download_docx_from_drive(args.drive_id, tmp.name)
    else:
        docx_path = args.docx_path
        if not os.path.exists(docx_path):
            sys.exit(f"Input file not found: {docx_path}")

    print(f"Parsing {docx_path}…")
    data = parse_docx(docx_path)
    print_stats(data)

    if args.json:
        with open(args.json, "w") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print(f"Wrote parsed JSON → {args.json}")

    if args.push:
        payloads = build_sheet_payloads(data)
        push_to_sheet(args.spreadsheet_id, payloads)
        print("Sheet updated. Now run `python3 scripts/parse_disciplines_to_json.py` "
              "to refresh data/disciplines.json.")
    else:
        print("\n--push not set — sheet was not modified.")


if __name__ == "__main__":
    main()
